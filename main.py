import json
import os
from abc import ABC, abstractmethod
from typing import TypeVar, Generic, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Ошибка: Библиотека python-docx не найдена. Установите её: pip install python-docx")
    exit()


class BaseStyle(ABC):
    def __init__(self, name: str):
        self.name = name

    @abstractmethod
    def to_dict(self) -> dict: pass

    @classmethod
    @abstractmethod
    def from_dict(cls, data: dict): pass

    def __eq__(self, other):
        if not isinstance(other, BaseStyle): return False
        return self.to_dict() == other.to_dict()


class DocxParagraphStyle(BaseStyle):
    def __init__(self, name: str,
                 font_name: str = "Calibri", font_size: float = 11,
                 bold: bool = False, italic: bool = False, color_rgb: Optional[list] = None,
                 alignment: int = 0, first_line_indent: float = 0.0,
                 left_indent: float = 0.0, space_after: float = 0.0):
        super().__init__(name)
        self.font_name = font_name
        self.font_size = font_size
        self.bold = bold
        self.italic = italic
        self.color_rgb = color_rgb
        self.alignment = alignment
        self.first_line_indent = first_line_indent
        self.left_indent = left_indent
        self.space_after = space_after

    def to_dict(self) -> dict:
        return {
            "type": "DocxParagraphStyle",
            "name": self.name,
            "font": {"name": self.font_name, "size": self.font_size, "bold": self.bold, "italic": self.italic,
                     "color": self.color_rgb},
            "paragraph": {"alignment": self.alignment, "first_line_indent": self.first_line_indent,
                          "left_indent": self.left_indent, "space_after": self.space_after}
        }

    @classmethod
    def from_dict(cls, data: dict):
        font = data.get("font", {})
        para = data.get("paragraph", {})
        return cls(
            name=data["name"],
            font_name=font.get("name", "Calibri"), font_size=font.get("size", 11),
            bold=font.get("bold", False), italic=font.get("italic", False), color_rgb=font.get("color", None),
            alignment=para.get("alignment", 0), first_line_indent=para.get("first_line_indent", 0.0),
            left_indent=para.get("left_indent", 0.0), space_after=para.get("space_after", 0.0)
        )

    def __str__(self):
        return f"Style '{self.name}': {self.font_name}, {self.font_size}pt"


class ApplyRule(ABC):
    def __init__(self, target_style_name: str):
        self.target_style_name = target_style_name

    @abstractmethod
    def match(self, paragraph_text: str) -> bool:
        pass


class UniversalRule(ApplyRule):
    def match(self, paragraph_text: str) -> bool:
        return True


class KeywordRule(ApplyRule):
    def __init__(self, target_style_name: str, keyword: str):
        super().__init__(target_style_name)
        self.keyword = keyword.lower()

    def match(self, paragraph_text: str) -> bool:
        return self.keyword in paragraph_text.lower()


class LengthRule(ApplyRule):
    def __init__(self, target_style_name: str, max_chars: int):
        super().__init__(target_style_name)
        self.max_chars = max_chars

    def match(self, paragraph_text: str) -> bool:
        t = paragraph_text.strip()
        return len(t) > 0 and len(t) <= self.max_chars


T = TypeVar('T', bound=BaseStyle)


class StyleCollection(Generic[T]):
    def __init__(self):
        self._items: List[T] = []

    def add(self, item: T):
        for idx, existing in enumerate(self._items):
            if existing.name == item.name:
                self._items[idx] = item
                return
        self._items.append(item)

    def __lshift__(self, item: T):
        self.add(item)
        return self

    def __getitem__(self, key: str) -> T:
        for item in self._items:
            if item.name == key: return item
        raise KeyError(f"Стиль {key} не найден")

    def count(self) -> int:
        return len(self._items)

    def get_all_items(self) -> List[T]:
        return self._items

    def contains(self, name: str) -> bool:
        return any(i.name == name for i in self._items)

    def clear(self):
        self._items.clear()

    def __eq__(self, other):
        if not isinstance(other, StyleCollection): return False
        return [i.to_dict() for i in self._items] == [i.to_dict() for i in other._items]

    def save_to_file(self, filepath: str):
        data = [item.to_dict() for item in self._items]
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"База стилей сохранена в {filepath}")

    def load_from_file(self, filepath: str):
        if not os.path.exists(filepath): return
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        self.clear()
        for entry in data:
            if entry.get("type") == "DocxParagraphStyle":
                self.add(DocxParagraphStyle.from_dict(entry))
        print(f"Загружено {len(self._items)} стилей")


class DocxManager:
    @staticmethod
    def extract_styles(docx_path: str) -> StyleCollection:
        collection = StyleCollection()
        if not os.path.exists(docx_path): return collection

        doc = Document(docx_path)
        print(f"Анализ файла {docx_path}...")

        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                if not style.font: continue

                font = style.font
                pf = style.paragraph_format

                rgb = [font.color.rgb[0], font.color.rgb[1], font.color.rgb[2]] if (
                            font.color and font.color.rgb) else None
                align = 0
                if pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    align = 1
                elif pf.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    align = 2
                elif pf.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    align = 3

                ds = DocxParagraphStyle(
                    name=style.name,
                    font_name=font.name if font.name else "Calibri",
                    font_size=font.size.pt if font.size else 11,
                    bold=bool(font.bold), italic=bool(font.italic), color_rgb=rgb,
                    alignment=align,
                    first_line_indent=pf.first_line_indent.cm if pf.first_line_indent else 0.0,
                    left_indent=pf.left_indent.cm if pf.left_indent else 0.0,
                    space_after=pf.space_after.pt if pf.space_after else 0.0
                )
                collection << ds
        return collection

    @staticmethod
    def update_style_definitions(doc, collection: StyleCollection):
        for my_style in collection.get_all_items():
            try:
                word_style = doc.styles[my_style.name]
            except KeyError:
                word_style = doc.styles.add_style(my_style.name, WD_STYLE_TYPE.PARAGRAPH)

            word_style.font.name = my_style.font_name
            if my_style.font_size: word_style.font.size = Pt(my_style.font_size)
            word_style.font.bold = my_style.bold
            word_style.font.italic = my_style.italic
            if my_style.color_rgb: word_style.font.color.rgb = RGBColor(*my_style.color_rgb)

            pf = word_style.paragraph_format
            if my_style.alignment == 0:
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif my_style.alignment == 1:
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif my_style.alignment == 2:
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif my_style.alignment == 3:
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            pf.first_line_indent = Cm(my_style.first_line_indent)
            pf.left_indent = Cm(my_style.left_indent)
            pf.space_after = Pt(my_style.space_after)

    @staticmethod
    def apply_rules_to_paragraphs(doc, rules: List[ApplyRule]):
        applied_count = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text: continue

            for rule in rules:
                if rule.match(text):
                    try:
                        p.style = doc.styles[rule.target_style_name]
                        applied_count += 1
                        break
                    except KeyError:
                        print(f"Внимание: Стиль '{rule.target_style_name}' назначен правилом, но отсутствует в файле.")

        print(f"Правила применены к {applied_count} параграфам.")

    @staticmethod
    def process_file(target_path: str, collection: StyleCollection, rules: List[ApplyRule] = None):
        if os.path.exists(target_path):
            doc = Document(target_path)
            print(f"Обработка файла: {target_path}")
        else:
            doc = Document()  # Новый файл
            print(f"Создан новый файл: {target_path}")

        DocxManager.update_style_definitions(doc, collection)

        if rules and len(rules) > 0:
            DocxManager.apply_rules_to_paragraphs(doc, rules)

        doc.save(target_path)
        print("Файл успешно сохранен.")


def edit_style_menu(style: DocxParagraphStyle):
    print(f"Редактируем {style.name}...")
    try:
        new_size = float(input(f"Размер ({style.font_size}): ") or style.font_size)
        style.font_size = new_size

        new_color = input(f"Цвет R G B (сейчас {style.color_rgb}): ")
        if new_color: style.color_rgb = [int(x) for x in new_color.split()]

        print("Стиль обновлен в памяти.")
    except Exception:
        print("Ошибка ввода")


def main():
    collection = StyleCollection()

    if not os.path.exists("demo.docx"):
        d = Document()
        d.add_paragraph("Введение")
        d.add_paragraph("Это очень короткий текст введения.")
        d.add_paragraph("Глава 1. Теория")
        d.add_paragraph("Здесь идет длинный текст, который должен быть обычным телом документа. " * 3)
        d.add_paragraph("Заключение")
        d.save("demo.docx")

    while True:
        print("\n=== DOCX Style Architect ===")
        print("1. Загрузить стили из файла (demo.docx)")
        print("2. Показать стили в памяти")
        print("3. Редактировать стиль")
        print("4. Сохранить/Загрузить JSON базы")
        print("--- Применение ---")
        print("5. Обновить определения стилей в файле (без изменения назначения)")
        print("6. ПРИМЕНИТЬ: Массово (один стиль ко всему)")
        print("7. ПРИМЕНИТЬ: С разделением (по правилам: заголовки, обычный текст)")
        print("0. Выход")

        choice = input("Выбор: ")

        if choice == "1":
            path = input("Файл (demo.docx): ") or "demo.docx"
            extracted = DocxManager.extract_styles(path)
            for item in extracted.get_all_items(): collection << item
            print(f"В памяти {collection.count()} стилей.")

        elif choice == "2":
            for s in collection.get_all_items(): print(s)

        elif choice == "3":
            name = input("Имя стиля: ")
            if collection.contains(name):
                edit_style_menu(collection[name])
            else:
                print("Нет такого стиля.")

        elif choice == "4":
            act = input("(s)ave / (l)oad: ")
            if act == 's':
                collection.save_to_file("db.json")
            elif act == 'l':
                collection.load_from_file("db.json")

        elif choice == "5":
            path = input("Файл (demo.docx): ") or "demo.docx"
            DocxManager.process_file(path, collection, rules=None)

        elif choice == "6":
            path = input("Файл (demo.docx): ") or "demo.docx"
            s_name = input("Имя стиля для всех абзацев (напр. Normal): ")
            if collection.contains(s_name):
                rules = [UniversalRule(s_name)]
                DocxManager.process_file(path, collection, rules)
            else:
                print("Такого стиля нет в коллекции. Сначала добавьте или загрузите его.")

        elif choice == "7":
            path = input("Файл (demo.docx): ") or "demo.docx"

            print("--- Формирование правил ---")
            print("Сценарий: Короткие строки -> Heading 1, Остальное -> Normal")

            h_style = input("Введите имя стиля заголовка (напр. Heading 1): ")
            b_style = input("Введите имя стиля текста (напр. Normal): ")

            if not (collection.contains(h_style) and collection.contains(b_style)):
                print("Ошибка: Указанные стили должны быть в загруженной коллекции!")
                continue

            rules = [
                LengthRule(target_style_name=h_style, max_chars=40),
                UniversalRule(target_style_name=b_style)
            ]

            DocxManager.process_file(path, collection, rules)

        elif choice == "0":
            break


if __name__ == "__main__":
    main()